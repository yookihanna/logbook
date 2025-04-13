import pandas as pd
from io import BytesIO
import streamlit as st
from openpyxl import Workbook
from docx import Document

# ========== Export to Excel ==========
def export_to_excel(df):
    output = BytesIO()
    writer = pd.ExcelWriter(output, engine='openpyxl')
    df.to_excel(writer, index=False, sheet_name='Logbook')
    writer.close()
    output.seek(0)

    st.download_button(
        label="📥 Download sebagai Excel",
        data=output,
        file_name='logbook.xlsx',
        mime='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet'
    )

# ========== Export to Word ==========
def export_to_word(df):
    document = Document()
    document.add_heading("Logbook Harian", 0)

    table = document.add_table(rows=1, cols=len(df.columns))
    table.style = 'Light Grid'
    hdr_cells = table.rows[0].cells
    for i, column in enumerate(df.columns):
        hdr_cells[i].text = column

    for index, row in df.iterrows():
        row_cells = table.add_row().cells
        for i, item in enumerate(row):
            row_cells[i].text = str(item)

    output = BytesIO()
    document.save(output)
    output.seek(0)

    st.download_button(
        label="📥 Download sebagai Word",
        data=output,
        file_name='logbook.docx',
        mime='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
    )
